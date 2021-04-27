import ast
import json

from django.http import FileResponse, HttpResponse
from django.shortcuts import render, redirect
from django.core.exceptions import ObjectDoesNotExist
from questionnaire import models
from docx import Document
# installation of both docx & python-docx
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Table
from reportlab.lib import colors
from django.db.models import Prefetch
import io
# from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
import os


def home(request):
    if request.session.is_empty() or (not request.session.get('id', False)):
        return render(request, 'home.html')
    else:
        cid = request.session['id']
        return render(request, 'home.html', {"id": cid})


def index(request):
    if request.session.is_empty() or (not request.session.get('id', False)):
        return render(request, 'index.html')
    else:
        cid = request.session['id']
        return render(request, 'index.html', {"id": cid})


def slr(request):
    if request.session.is_empty() or (not request.session.get('id', False)):
        return render(request, 'SLR.html')
    else:
        cid = request.session['id']
        questions = models.QuestionSLR.objects.prefetch_related(
            Prefetch("answer", queryset=models.Answer.objects.filter(cid=cid))
        ).all()
        q: models.QuestionSLR
        return render(request, 'SLR.html', {"id": cid, "questions": questions})


def baseline(request):
    if request.session.is_empty() or (not request.session.get('id', False)):
        return render(request, 'tolerance.html')
    else:
        cid = request.session['id']
        questions = models.QuestionTolerance.objects.all()
        return render(request, 'tolerance.html', {"id": cid, "questions": questions})


def invtype(request):
    if request.session.is_empty() or (not request.session.get('id', False)):
        return render(request, 'invtype.html')
    else:
        cid = request.session['id']
        questions = models.QuestionInvtype.objects.all()
        """
        paginator = Paginator(questions, 2)
        page = request.GET.get('page', 1)
        currPage = int(page)
        try:
            page_list = paginator.page(page)
        except PageNotAnInteger:
            page_list = paginator.page(1)
        except EmptyPage:
            page_list = paginator.page(paginator.num_pages)
        """

        # return render(request, 'invtype.html', locals())
        return render(request, 'invtype.html', {"id": cid, "questions": questions})


def esg(request):
    if request.session.is_empty() or (not request.session.get('id', False)):
        return render(request, 'ESG.html')
    else:
        cid = request.session['id']
        questions = models.QuestionESG.objects.all()
        return render(request, 'ESG.html', {"id": cid, "questions": questions})


def trends(request):
    if request.session.is_empty() or (not request.session.get('id', False)):
        return render(request, 'trends.html')
    else:
        cid = request.session['id']
        questions = models.QuestionTrends.objects.all()
        return render(request, 'trends.html', {"id": cid, 'questions': questions})


def clientLogin(request):
    if request.method == 'GET':
        return render(request, 'login.html')
    elif request.method == 'POST':
        cid = request.POST.get('id')
        password = request.POST.get('password')

        try:
            client = models.Client.objects.get(id=cid)

            if password == client.password:
                request.session['id'] = cid
                return render(request, 'home.html', {'id': cid})
            else:
                return render(request, 'login.html', {'error': 'incorrect id or password'})
        except ObjectDoesNotExist:
            return render(request, 'login.html', {'error': 'incorrect id or password'})


def clientLogout(request):
    request.session.flush()
    return render(request, 'home.html')


def submitTolerance(request):
    cid = request.POST.get("cid") or request.GET.get("cid")
    nextQ = models.QuestionInvtype.objects.prefetch_related(
        Prefetch("answer", queryset=models.Answer.objects.filter(cid=cid))
    ).all()
    if request.method == 'POST':
        try:
            pid = request.POST.get('pid')
            question = models.QuestionTolerance.objects.all()

            client = models.Client.objects.get(id=cid)
            for q in question:
                qid = str(q.id)
                a = request.POST.get(qid)
                if models.Answer.objects.filter(cid=client, pid=pid, qid=qid):
                    models.Answer.objects.filter(cid=client, pid=pid, qid=qid).update(answer=a, question_tolerance=q)
                else:
                    models.Answer.objects.create(cid=client, pid=pid, qid=qid, answer=a, question_tolerance=q)
            if request.POST.get("get_report") is not None:
                return redirect("/submitTrends")
            return render(request, 'invtype.html', {'id': cid, 'questions': nextQ})
        except:
            if request.POST.get("get_report") is not None:
                return redirect("/submitTrends")
            return render(request, 'invtype.html', {'id': cid, 'questions': nextQ})
    else:
        return render(request, 'invtype.html', {'id': cid, 'questions': nextQ})


"""
def submitInv(request):
    if request.method == 'POST':
        cid = request.POST.get('cid')
        pid = request.POST.get('pid')
        client = models.Client.objects.get(id=cid)
        # question = models.QuestionInvtype.objects.all()
        page = request.GET.get('page', 1)
        nextPage = request.GET.get('next', "")
        questions = request.POST.get('qid')
        for q in questions:
            qid = str(q.id)
            a = request.POST.get(qid)
            models.Answer.objects.create(cid=client, pid=pid, qid=qid, answer=a)
        return render(request, 'invtype.html', locals())
"""


def submitInvtype(request):
    cid = request.POST.get('cid') or request.GET.get("cid")
    nextQ = models.QuestionESG.objects.prefetch_related(
        Prefetch("answer", queryset=models.Answer.objects.filter(cid=cid))
    ).all()
    if request.method == 'POST':
        try:
            pid = request.POST.get('pid')
            question = models.QuestionInvtype.objects.all()
            client = models.Client.objects.get(id=cid)
            counta = 0
            countb = 0
            countc = 0
            countd = 0

            for q in question:
                qid = str(q.id)
                a = request.POST.get(qid)
                if models.Answer.objects.filter(cid=client, question_inv_type=q):
                    models.Answer.objects.filter(cid=client, qid=qid, question_inv_type=q).update(
                        answer=a,
                        question_inv_type=q
                    )
                else:
                    models.Answer.objects.create(cid=client, pid=pid, answer=a, question_inv_type=q)

                if a == 'A':
                    counta += 1
                if a == 'B':
                    countb += 1
                if a == 'C':
                    countc += 1
                if a == 'D':
                    countd += 1

            if counta in range(5, 11):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Preserver')
                else:
                    models.Result.objects.create(cid=client, result='Preserver')
            elif countb in range(5, 11):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Accumulator')
                else:
                    models.Result.objects.create(cid=client, result='Accumulator')
            elif countc in range(5, 11):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Independent')
                else:
                    models.Result.objects.create(cid=client, result='Independent')
            elif countd in range(5, 11):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Follower')
                else:
                    models.Result.objects.create(cid=client, result='Follower')
            elif counta == 4 and countb in (3, 4):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Preserver+Accumulator')
                else:
                    models.Result.objects.create(cid=client, result='Preserver+Accumulator')
            elif counta == 4 and countc in (3, 4):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Preserver+Independent')
                else:
                    models.Result.objects.create(cid=client, result='Preserver+Independent')
            elif counta == 4 and countd in (3, 4):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Preserver+Follower')
                else:
                    models.Result.objects.create(cid=client, result='Preserver+Follower')
            elif countb == 4 and countc in (3, 4):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Independent+Accumulator')
                else:
                    models.Result.objects.create(cid=client, result='Independent+Accumulator')
            elif countb == 4 and countd in (3, 4):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Follower+Accumulator')
                else:
                    models.Result.objects.create(cid=client, result='Follower+Accumulator')
            elif countc == 4 and countd in (3, 4):
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='Follower+Independent')
                else:
                    models.Result.objects.create(cid=client, result='Follower+Independent')
            else:
                if models.Result.objects.filter(cid=client):
                    models.Result.objects.filter(cid=client).update(result='No definite investor type')
                else:
                    models.Result.objects.create(cid=client, result='No definite investor type')

            if request.POST.get("get_report") is not None:
                return redirect("/submitTrends")
            return render(request, 'ESG.html', {'id': cid, 'questions': nextQ})
        except:
            if request.POST.get("get_report") is not None:
                return redirect("/submitTrends")
            return render(request, 'ESG.html', {'id': cid, 'questions': nextQ})
    else:
        return render(request, 'ESG.html', {'id': cid, 'questions': nextQ})


def submitESG(request):
    cid = request.POST.get('cid') or request.GET.get('cid')
    nextQ = models.QuestionTrends.objects.prefetch_related(
        Prefetch("answer", queryset=models.Answer.objects.filter(cid=cid))
    ).all()
    if request.method == 'POST':
        try:
            pid = request.POST.get('pid')
            question = models.QuestionESG.objects.all()
            client = models.Client.objects.get(id=cid)
            for q in question:
                qid = str(q.id)
                a = request.POST.getlist(qid)
                if 'null' in a:
                    a.remove('null')
                if models.Answer.objects.filter(cid=client, pid=pid, qid=qid):
                    models.Answer.objects.filter(cid=client, pid=pid, qid=qid).update(answer=a, question_esg=q)
                else:
                    models.Answer.objects.create(cid=client, pid=pid, qid=qid, answer=a, question_esg=q)
            if request.POST.get("get_report") is not None:
                return redirect("/submitTrends")
            return render(request, 'trends.html', {'id': cid, 'questions': nextQ})
        except:
            if request.POST.get("get_report") is not None:
                return redirect("/submitTrends")
            return render(request, 'trends.html', {'id': cid, 'questions': nextQ})
    else:
        return render(request, 'trends.html', {'id': cid, 'questions': nextQ})


def submitSLR(request):
    cid = request.POST.get('cid') or request.GET.get('cid')
    nextQ = models.QuestionTolerance.objects.prefetch_related(
        Prefetch("answer", queryset=models.Answer.objects.filter(cid=cid))
    ).all()
    if request.method == 'POST':
        try:
            pid = request.POST.get('pid')
            question = models.QuestionSLR.objects.all()

            client = models.Client.objects.get(id=cid)
            for q in question:
                qid = str(q.id)
                a = request.POST.get(qid, None)
                b = request.POST.get(qid + "_text")
                if b and not a:
                    a = b
                if models.Answer.objects.filter(cid=client, pid=pid, qid=qid) and a:
                    models.Answer.objects.filter(cid=client, pid=pid, qid=qid).update(
                        answer=a,
                        question_slr=q.id,
                        textbox=b,
                    )
                else:
                    if a:
                        models.Answer.objects.create(
                            cid=client,
                            pid=pid,
                            qid=qid,
                            answer=a,
                            question_slr=q,
                            textbox=b,
                        )
            # SLR calculation
            # 1. original SLR
            income = models.Answer.objects.filter(cid=client, pid=pid, qid=8)[0].answer
            expense = models.Answer.objects.filter(cid=client, pid=pid, qid=11)[0].answer
            otherAsset = models.Answer.objects.filter(cid=client, pid=pid, qid=6)[0].answer

            income_growth = models.Discount.objects.all()
            if income_growth:
                income_growth = income_growth[0].income_growth
            else:
                income_growth = 2

            expense_growth = models.Discount.objects.all()
            if expense_growth:
                expense_growth = expense_growth[0].expense_growth
            else:
                expense_growth = 1

            discount = models.Discount.objects.all()
            if discount:
                discount = discount[0].discounting_factor
            else:
                discount = 2

            year = models.Answer.objects.filter(cid=client, pid=pid, qid=3)[0].answer

            # print(income_growth)
            # # print(year)
            # print(discount)
            # print(expense)

            # print("-----")
            # print((1-((1+expense_growth)/(1+discount))))

            th_final = float(1 - (1 + expense_growth) / (1 + discount)) + float(otherAsset)
            th_1 = (1 - ((1 + income_growth) / (1 + discount)) ** (float(year) + 1))
            th_2 = (1 - ((1 + expense_growth) / (1 + discount)) ** (float(year) + 1))
            # print(1-((1+income_growth)/(1+discount)))

            # print(th_1)
            # print(th_2)
            # print(th_final)
            asset = float(income) * th_1 / float(1 - (1 + income_growth) / (1 + discount)) - float(
                expense) * th_2 / th_final

            liability = models.Answer.objects.filter(cid=client, pid=pid, qid=10)[0].answer
            SLR = float((asset - float(liability)) / asset)
            # 2. additional SLR
            otherSLR = 0
            age = models.Answer.objects.filter(cid=client, pid=pid, qid=2)[0].answer
            if age == 'B':
                otherSLR += 1
            elif age == 'C':
                otherSLR += 4
            elif age == 'D':
                otherSLR += 6
            elif age == 'E':
                otherSLR += 9
            education = models.Answer.objects.filter(cid=client, pid=pid, qid=4)[0].answer
            if education == 'B':
                otherSLR += 1
            elif education == 'C':
                otherSLR += 2
            elif education == 'D':
                otherSLR += 3
            elif education == 'E':
                otherSLR += 4
            dependents = models.Answer.objects.filter(cid=client, pid=pid, qid=5)[0].answer
            if dependents == 'B':
                otherSLR += 1
            elif dependents == 'C':
                otherSLR += 2
            elif dependents == 'D':
                otherSLR += 3
            elif dependents == 'E':
                otherSLR += 4
            retirement = models.Answer.objects.filter(cid=client, pid=pid, qid=9)[0].answer
            if retirement == 'B':
                otherSLR += 0.5
            elif retirement == 'C':
                otherSLR += 1
            elif retirement == 'D':
                otherSLR += 3
            elif retirement == 'E':
                otherSLR += 5.5
            newJob = models.Answer.objects.filter(cid=client, pid=pid, qid=13)[0].answer
            if newJob == 'B':
                otherSLR += 2
            elif newJob == 'C':
                otherSLR += 3
            elif newJob == 'D':
                otherSLR += 5
            payStructure = models.Answer.objects.filter(cid=client, pid=pid, qid=14)[0].answer
            if payStructure == 'B':
                otherSLR += 1
            elif payStructure == 'C':
                otherSLR += 2
            elif payStructure == 'D':
                otherSLR += 3
            elif payStructure == 'E':
                otherSLR += 4
            knowledge = models.Answer.objects.filter(cid=client, pid=pid, qid=15)[0].answer
            if knowledge == 'B':
                otherSLR += 1
            elif knowledge == 'C':
                otherSLR += 3
            elif knowledge == 'D':
                otherSLR += 6
            # final calculation
            resultSLR = float(SLR * 0.8 + otherSLR / 80 * 0.2)
            if models.ResultSLR.objects.filter(cid=client):
                models.ResultSLR.objects.filter(cid=client).update(slr=resultSLR)
            else:
                models.ResultSLR.objects.create(cid=client, slr=resultSLR)

            if request.POST.get("get_report") is not None:
                return redirect("/submitTrends")
            return render(request, 'tolerance.html', {'id': cid, 'questions': nextQ})
        except Exception as ex:
            print(ex)
            if request.POST.get("get_report") is not None:
                return redirect("/submitTrends")
            return render(request, 'tolerance.html', {'id': cid, 'questions': nextQ})
    else:
        cid = request.GET.get('cid')
        return render(request, 'tolerance.html', {'id': cid, 'questions': nextQ})


def submitTrends(request):
    cid = request.session['id'] or request.POST.get('cid')
    if request.method == 'POST':
        cid = request.POST.get('cid')
        pid = request.POST.get('pid')
        question = models.QuestionTrends.objects.all()
        client = models.Client.objects.get(id=cid)
        for q in question:
            qid = str(q.id)
            a = request.POST.getlist(qid)
            if 'null' in a:
                a.remove('null')
            if models.Answer.objects.filter(cid=client, pid=pid, qid=qid):
                models.Answer.objects.filter(cid=client, pid=pid, qid=qid).update(answer=a, question_trend=q)
            else:
                models.Answer.objects.create(cid=client, pid=pid, qid=qid, answer=a, question_trend=q)
    SLR = models.ResultSLR.objects.filter(cid=cid)

    if SLR:
        SLR = SLR[0].slr

    BIT = models.Result.objects.filter(cid=cid)
    temp = 'a'
    if BIT:
        BIT = BIT[0].result

        if BIT[0] == 'A' or BIT[0] == 'I':
            temp = 'an'
    else:
        BIT = ""

    description1 = ""
    file = Document(str(os.getcwd()) + "/statics/BIT.docx")
    for i in range(len(file.paragraphs)):
        if file.paragraphs[i].text == ("Investor Type: " + str.upper(str(BIT))):
            description1 = file.paragraphs[i + 1].text
            for j in range(2, 5):
                if file.paragraphs[i + j].text[0:14] == 'Investor Type:':
                    break
                else:
                    description1 += ('\n\n' + file.paragraphs[i + j].text)
            break
    description2 = '\n\n' + file.paragraphs[-3].text + '\n\n' + file.paragraphs[-2].text + '\n\n' + file.paragraphs[
        -1].text
    description = description1 + description2

    esg = models.Answer.objects.filter(cid=cid, pid='ESG')
    if esg:
        esg = esg[0].answer
    ESG = ""
    for answer in esg.strip('[').strip(']').split(','):
        if '' != answer:
            e = models.QuestionESG.objects.filter(id=1).values_list('option' + answer.strip(' ').strip("'"), flat=True)

            if e:
                e = e[0]
            else:
                e = ""
            if e == 'I do not take consideration ESG factors':
                ESG += ", " + "No ESG Preferences"
            else:
                ESG += ", " + e
    ESG = ESG[2:]
    trends = models.Answer.objects.filter(cid=cid, pid='Trends').values_list('qid', 'answer')
    TRENDS = ""
    for (question, answers) in trends:
        for answer in answers.strip('[').strip(']').split(','):
            if '' != answer:
                t = models.QuestionTrends.objects.filter(id=question).values_list(
                    'option' + answer.strip('').strip(' ').strip("'"), flat=True)[0]
                if t not in ['Other', 'None of These', 'I do not have a view']:
                    TRENDS += ", " + t
    TRENDS = TRENDS[2:]
    if not ESG:
        ESG = 'We have investment strategies that are customized to your ESG preferences.  ' \
              'Please contact us to discuss how we implement these strategies.'
    description1 = 'You have a financially risky lifestyle.  ' \
                   'Depending upon specifics, you may have the propensity ' \
                   'to outspend your asset base.  You likely need a financial plan and ' \
                   'portfolio review to make sure you are on track to reach your financial goals.'
    description2 = 'You have a moderately risky financial lifestyle.  ' \
                   'Depending upon specifics, you likely are on track to ' \
                   'reach your financial goals.  With that said, you may need a ' \
                   'financial plan and portfolio review to make sure that you will not ' \
                   'outspend your asset base.'
    description3 = 'You have a low-risk financial lifestyle.  ' \
                   'Depending upon specifics, you likely are on track to reach ' \
                   ' financial goals.  With that said, you may need a financial plan ' \
                   'and portfolio review to determine your exact financial priorities.'
    return render(request, 'result.html',
                  {'id': cid, 'slr': SLR, 'temp': temp, 'bit': BIT, 'description': description, 'esg': ESG,
                   'trends': TRENDS, 'description1': description1, 'description2': description2, 'description3': description3})



def result(request):
    if request.session.is_empty():
        return render(request, 'result.html')
    else:
        try:
            cid = request.session['id']
            SLR = models.ResultSLR.objects.filter(cid=cid)[0].slr
            BIT = models.Result.objects.filter(cid=cid)[0].result
            if BIT[0] == 'A' or BIT[0] == 'I':
                temp = 'an'
            else:
                temp = 'a'
            file = Document(str(os.getcwd()) + "\\statics\\BIT.docx")
            for i in range(len(file.paragraphs)):
                if file.paragraphs[i].text == ("Investor Type: " + str.upper(str(BIT))):
                    description1 = file.paragraphs[i + 1].text
                    for j in range(2, 5):
                        if file.paragraphs[i + j].text[0:14] == 'Investor Type:':
                            break
                        else:
                            description1 += ('\n\n' + file.paragraphs[i + j].text)
                    break
            description2 = '\n\n' + file.paragraphs[-3].text + '\n\n' + file.paragraphs[-2].text + '\n\n' + \
                           file.paragraphs[-1].text
            description = description1 + description2
            esg = models.Answer.objects.filter(cid=cid, pid='ESG', qid=1)[0].answer
            ESG = ""
            for answer in esg.strip('[').strip(']').split(','):
                e = \
                    models.QuestionESG.objects.filter(id=1).values_list('option' + answer.strip(' ').strip("'"),
                                                                        flat=True)[
                        0]
                if e == 'I do not take consideration ESG factors':
                    ESG += ", " + "No ESG Preferences"
                else:
                    ESG += ", " + e
            ESG = ESG[2:]
            trends = models.Answer.objects.filter(cid=cid, pid='Trends').values_list('qid', 'answer')
            TRENDS = ""
            for (question, answers) in trends:
                for answer in answers.strip('[').strip(']').split(','):
                    t = models.QuestionTrends.objects.filter(id=question).values_list(
                        'option' + answer.strip(' ').strip("'"), flat=True)[0]
                    if t not in ['Other', 'None of These', 'I do not have a view']:
                        TRENDS += ", " + t
            TRENDS = TRENDS[2:]
            return render(request, 'result.html',
                          {"id": cid, 'slr': SLR, 'temp': temp, 'bit': BIT, 'description': description, 'esg': ESG,
                           'trends': TRENDS})
        except:
            return render(request, 'result.html')


def dlreport(request):
    cid = request.session['id']

    SLR = models.ResultSLR.objects.filter(cid=cid)

    if SLR:
        SLR = SLR[0].slr

    BIT = models.Result.objects.filter(cid=cid)[0].result
    file = Document(str(os.getcwd()) + "/statics/BIT.docx")
    for i in range(len(file.paragraphs)):
        if file.paragraphs[i].text == ("Investor Type: " + str.upper(str(BIT))):
            description1 = file.paragraphs[i + 1].text
            for j in range(2, 5):
                if file.paragraphs[i + j].text[0:14] == 'Investor Type:':
                    break
                else:
                    description1 += ('\n' + file.paragraphs[i + j].text)
            break
    description2 = '\n' + file.paragraphs[-3].text + '\n' + file.paragraphs[-2].text + '\n' + file.paragraphs[-1].text
    esg = models.Answer.objects.filter(cid=cid, pid='ESG', qid=1)

    if esg:
        esg = esg[0].answer
    else:
        esg = ""
    ESG = ""
    for answer in esg.strip('[').strip(']').split(','):
        if '' != answer:
            e = models.QuestionESG.objects.filter(id=1).values_list('option' + answer.strip(' ').strip("'"), flat=True)

            if e:
                e = e[0]
            if e == 'I do not take consideration ESG factors':
                ESG += ", " + "No ESG Preferences"
            else:
                ESG += ", " + e
    ESG = ESG[2:]
    trends = models.Answer.objects.filter(cid=cid, pid='Trends').values_list('qid', 'answer')
    TRENDS = ""
    for (question, answers) in trends:
        for answer in answers.strip('[').strip(']').split(','):
            if '' != answer:
                t = \
                    models.QuestionTrends.objects.filter(id=question).values_list(
                        'option' + answer.strip(' ').strip("'"),
                        flat=True)[0]
                if t not in ['Other', 'None of These', 'I do not have a view']:
                    TRENDS += ", " + t
    TRENDS = TRENDS[2:]

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    # p.drawImage(str(os.getcwd())+"/statics/images/logo-blue.png", width-100, height-50, 72, 30.82)
    p.setFont('Times-Roman', 16)
    if BIT == 'No definite investor type':
        title = "Your behavioral investor type is undefined"
        p.drawCentredString(width / 2, height - 70, title)
    elif BIT[0] == 'A' or BIT[0] == 'I':
        title1 = "You are an"
        title2 = str.upper(BIT)
        p.drawCentredString(width / 2 - 45, height - 70, title1)
        t = p.beginText()
        t.setTextOrigin(width / 2 - 7, height - 70)
        t.setFont('Times-Bold', 16)
        t.textLine(title2)
        p.drawText(t)
    else:
        title1 = "You are a"
        title2 = str.upper(BIT)
        p.drawCentredString(width / 2 - 45, height - 70, title1)
        t = p.beginText()
        t.setTextOrigin(width / 2 - 7, height - 70)
        t.setFont('Times-Bold', 16)
        t.textLine(title2)
        p.drawText(t)
    q1 = "Your SLR score: "
    if not SLR:
        SLR = 0
    a1 = simpleSplit("%.4f" % (float(SLR) * 100) + "%", 'Times-Bold', 14, maxWidth=500)
    q2 = "ESG preferences: "
    a2 = simpleSplit(ESG, 'Times-Bold', 14, maxWidth=500)
    q3 = "You believe in trends including: "
    a3 = simpleSplit(TRENDS, 'Times-Bold', 14, maxWidth=500)
    q4 = "Your behavioral investor type: "
    a4 = simpleSplit(BIT, 'Times-Bold', 14, maxWidth=500)
    a5 = simpleSplit(description1, 'Times-Roman', 14, maxWidth=500)
    a6 = simpleSplit(description2, 'Times-Roman', 14, maxWidth=500)

    text = p.beginText()
    text.setTextOrigin(55, height - 120)
    text.setFont('Times-Roman', 15, leading=1.5 * p._leading)
    text.textLines(q4)
    text.setFont('Times-Bold', 14, leading=1.5 * p._leading)
    text.textLines(a4)
    p.drawText(text)

    for line in a5:
        _, y = text.getCursor()
        if y < 50:
            p.drawText(text)
            p.showPage()
            p.setFont('Times-Roman', 16)
            text = p.beginText()
            text.setTextOrigin(55, height - 70)
        text.setFont('Times-Roman', 14, leading=1.5 * p._leading)
        text.textLines(line)
    text.textLines(' ')
    for line in a6:
        _, y = text.getCursor()
        if y < 50:
            p.drawText(text)
            p.showPage()
            p.setFont('Times-Italic', 16)
            text = p.beginText()
            text.setTextOrigin(55, height - 70)
        text.setFont('Times-Italic', 14, leading=1.5 * p._leading)
        text.textLines(line)

    text = p.beginText()
    text.setTextOrigin(55, height - 220)
    text.setFont('Times-Roman', 15, leading=1.5 * p._leading)
    text.textLines(q1)
    text.setFont('Times-Bold', 14, leading=1.5 * p._leading)
    text.textLines(a1)
    p.drawText(text)
    data = [
        ['SLR', 'Below 40%', '40% - 70%', 'Above 70%'],
        ['Lifestyle Description', 'Risky Lifestyle', 'Moderate Risk Lifestyle', 'Low Risk Lifestyle']
    ]
    table = Table(data, colWidths=[130, 100, 145, 120], rowHeights=2 * [25], style=[
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        # ('SPAN', (0, 0), (0, 1)),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTSIZE', (0, 0), (-1, -1), 14),
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
    ])
    table.wrap(500, 50)
    table.drawOn(p, 55, height - 315)
    text = p.beginText()
    text.setTextOrigin(55, height - 345)
    temp = dict(q2=q2, a2=a2, q3=q3, a3=a3)
    for i in range(len(temp)):
        if list(temp.keys())[i][0] == 'q':
            text.setFont('Times-Roman', 15, leading=1.5 * p._leading)
        else:
            text.setFont('Times-Bold', 14, leading=1.5 * p._leading)
        text.textLines(list(temp.values())[i])
    # for line in a5:
    #     _, y = text.getCursor()
    #     if y < 50:
    #         p.drawText(text)
    #         p.showPage()
    #         p.setFont('Times-Roman', 16)
    #         text = p.beginText()
    #         text.setTextOrigin(55, height - 70)
    #     text.setFont('Times-Roman', 14, leading=1.5 * p._leading)
    #     text.textLines(line)
    # text.textLines(' ')
    # for line in a6:
    #     _, y = text.getCursor()
    #     if y < 50:
    #         p.drawText(text)
    #         p.showPage()
    #         p.setFont('Times-Italic', 16)
    #         text = p.beginText()
    #         text.setTextOrigin(55, height - 70)
    #     text.setFont('Times-Italic', 14, leading=1.5 * p._leading)
    #     text.textLines(line)
    p.drawText(text)
    p.showPage()
    p.save()
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename='Report.pdf')
